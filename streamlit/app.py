import streamlit as st
import streamlit.components.v1 as components
from streamlit_option_menu import option_menu

import base64
from PyPDF2 import PdfReader

from io import BytesIO # 이건 뭐지
from docx import Document # 다운받아야함
from docx2pdf import convert
from fpdf import FPDF
import tempfile
# import pyhwp  # 다운받아야함 # 얘는 호환이 안 되나봄. 일단 파일은 docx, pdf만 해보겠음

### 아래는 제가 모듈화 한것(여진)
# import connect_db
# import get_sop 

### 샘플 데이터 ###
docs = [
    {"title": "Guidance for Industry: Manufacturing Considerations for Licensed and Investigational Cellular and Gene Therapy Products", "changes": 7, "content": """1. Quality Control Focus
Guidance for Industry: Quality and Manufacturing Controls for Cellular and Gene Therapy Products
Quality control programs must ensure product identity, purity, and potency throughout manufacturing. ...

2. Raw Materials Focus
Guidance for Industry: Source Materials and Manufacturing Considerations in Cellular and Gene Therapy
Source materials such as donor cells, viral vectors, and ancillary reagents must be well-characterized..."""},

    {"title": "ENFORCEMENT RULE OF THE DECREE ON FACILITY STANDARDS FOR MEDICINAL PRODUCT MANUFACTURERS AND IMPORTERS", "changes": 3, "content": "ENFORCEMENT RULE 내용..."},

    {"title": "시진핑의 신-10개년 약 조제 규정", "changes": 2, "content": "중국 규정 상세 내용..."},

    {"title": "약 배달원의 항의로 고쳐진 아무개 규정", "changes": 4, "content": "배달원 규정 상세 내용..."},

    {"title": "Translating GMP Standards into Harmonized Practices for Advanced Therapies", "changes": 1, "content": "국제 조화 규정 내용..."},
]

### 환경 세팅 ###
# db연결
# session_state 설정

if 'uploaded_sop' not in st.session_state:
    st.session_state['uploaded_sop'] = []

st.session_state.uploaded_sop = [
    {'document_name':''},{'document_content':''}
    ]

# 세션 상태 저장
if "selected" not in st.session_state:
    st.session_state["selected"] = None

### 화면 구성 ###
st.set_page_config(
    page_title="GMP AUTO COMPLIANCE", ##NGYEO변경부분(1차)
    layout="wide",  # 화면을 가로로 넓게
    initial_sidebar_state="expanded"  # 사이드바를 기본으로 펼치기
)

st.title('GMP AUTO COMPLIANCE') # 제목 부분

# 세션 상태 초기화
if "what" not in st.session_state:
    st.session_state.what = None

with st.sidebar:
    st.title('GMP AUTO COMPLIANCE')
    choice = option_menu("Menu", ["SOP 업로드", "최근 GMP 변경사항", "SOP 수정사항", "과거 기록"],
                         icons=['bi bi-file-earmark-arrow-down', 'bi bi-bell', 'bi bi-robot', 'bi bi-clock-history'],
                         menu_icon="app-indicator", default_index=0,
                         styles={
        "container": {"padding": "4!important", "background-color": "#fafafa"},
        "icon": {"color": "black", "font-size": "21px"},
        "nav-link": {"font-size": "16px", "text-align": "left", "margin":"0px", "--hover-color": "#fafafa"},
        "nav-link-selected": {"background-color": "#08c7b4"},
    }
    )

# 선택에 따른 화면 분기
if choice == "SOP 업로드":
    st.session_state.what = "SOP 업로드"
    uploaded_sop = st.file_uploader('여기 SOP 파일을 업로드 해주세요', type=['docx','pdf'])

    if uploaded_sop is not None:
        file_type = uploaded_sop.name.split(".")[-1].lower()

        if file_type == "docx":
            doc = Document(uploaded_sop)
            content = "\n".join([p.text for p in doc.paragraphs])
            st.write("DOCX 내용", content, height=600)

            ####################
            # 1. DOCX → PDF 변환
            doc = Document(uploaded_sop)
            pdf = FPDF()
            pdf.add_page()
            # 한글 폰트 추가 (나눔고딕 또는 원하는 TTF 폰트 경로 설정)
            pdf.add_font('NanumGothic', '', 'NanumGothic.ttf', uni=True)
            pdf.set_font('NanumGothic', '', 14)
            pdf.set_auto_page_break(auto=True, margin=15)
            pdf.set_font("Arial", size=12)
            
            for para in doc.paragraphs:
                pdf.multi_cell(0, 10, para.text)
            
            tmp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")
            pdf.output(tmp_pdf.name)
            pdf_path = tmp_pdf.name

                    # 2. PDF → Base64
            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()
            b64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            
            # 3. PDF.js 렌더링
            pdfjs_html = f"""
            <html>
            <head>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.16.105/pdf.min.js"></script>
                <style>
                    body {{ margin: 0; }}
                    #pdf-viewer {{ width: 100%; height: 100vh; overflow: auto; }}
                    canvas {{ display: block; margin: auto; }}
                </style>
            </head>
            <body>
                <div id="pdf-viewer"></div>
                <script>
                    const pdfData = atob("{b64_pdf}");
                    const pdfjsLib = window['pdfjs-dist/build/pdf'];
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.16.105/pdf.worker.min.js';

                    pdfjsLib.getDocument({{data: pdfData}}).promise.then(pdf => {{
                        const viewer = document.getElementById('pdf-viewer');
                        for (let i = 1; i <= pdf.numPages; i++) {{
                            pdf.getPage(i).then(page => {{
                                const canvas = document.createElement('canvas');
                                viewer.appendChild(canvas);
                                const context = canvas.getContext('2d');
                                const viewport = page.getViewport({{ scale: 1.5 }});
                                canvas.height = viewport.height;
                                canvas.width = viewport.width;
                                page.render({{ canvasContext: context, viewport: viewport }});
                            }});
                        }}
                    }});
                </script>
            </body>
            </html>
            """
            
            st.components.v1.html(pdfjs_html, height=800, scrolling=True)


        elif file_type == "pdf":
            pdf_bytes = uploaded_sop.read()
            b64_pdf = base64.b64encode(pdf_bytes).decode('utf-8')
            
            pdfjs_html = f"""
            <html>
            <head>
                <script src="https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.16.105/pdf.min.js"></script>
                <style>
                    body {{ margin: 0; }}
                    #pdf-viewer {{ width: 100%; height: 100vh; overflow: auto; }}
                    canvas {{ display: block; margin: auto; }}
                </style>
            </head>
            <body>
                <div id="pdf-viewer"></div>
                <script>
                    const pdfData = atob("{b64_pdf}");
                    const pdfjsLib = window['pdfjs-dist/build/pdf'];
                    pdfjsLib.GlobalWorkerOptions.workerSrc = 'https://cdnjs.cloudflare.com/ajax/libs/pdf.js/2.16.105/pdf.worker.min.js';

                    pdfjsLib.getDocument({{data: pdfData}}).promise.then(pdf => {{
                        const viewer = document.getElementById('pdf-viewer');
                        for (let i = 1; i <= pdf.numPages; i++) {{
                            pdf.getPage(i).then(page => {{
                                const canvas = document.createElement('canvas');
                                viewer.appendChild(canvas);
                                const context = canvas.getContext('2d');
                                const viewport = page.getViewport({{ scale: 1.5 }});
                                canvas.height = viewport.height;
                                canvas.width = viewport.width;
                                page.render({{ canvasContext: context, viewport: viewport }});
                            }});
                        }}
                    }});
                </script>
            </body>
            </html>
            """
            
            st.components.v1.html(pdfjs_html, height=800, scrolling=True)
        
elif choice == "최근 GMP 변경사항":
    st.session_state.what = "최근 GMP 변경사항"

    st.info('🔄마지막 업데이트: 2025-09-23 14:52:14')  # 파란색 영역

    col1, col2 = st.columns(2)
    col1.write("왼쪽")
    col2.write("오른쪽")

    # 레이아웃
    col1, col2 = st.columns([2,2])

    # 왼쪽: 스크롤 가능한 문서 목록
    with col1:
        st.subheader("[변경, 추가된 GMP]")
        with st.container(height=500):  # 스크롤 영역
            for i, doc in enumerate(docs):
                bg = "#d7f2ff" if st.session_state["selected"] == i else "#f9f9f9"
                with st.container(border=True):
                    st.markdown(
                        f"""
                        <div style="background-color:{bg}; padding:10px; border-radius:10px; margin-bottom:10px;">
                            <b>GMP 제목:</b> {doc["title"]}<br>
                            <b>변경된 부분 개수:</b> {doc["changes"]}
                        </div>
                        """,
                        unsafe_allow_html=True
                    )
                    if st.button("👁 보기", key=f"btn_{i}"):
                        st.session_state["selected"] = i

    # 오른쪽: 상세 내용
    with col2:
        if st.session_state["selected"] is not None:
            st.subheader(docs[st.session_state["selected"]]["title"])
            st.text_area("내용", docs[st.session_state["selected"]]["content"], height=600)
        else:
            st.info("왼쪽에서 GMP 문서를 선택하세요.")

elif choice == "SOP 수정사항":
    st.session_state.what = "SOP 수정사항"
    if st.session_state.uploaded_sop != []:
        st.write('전문')

    uploaded_file = st.file_uploader("문서 업로드 (DOCX, PDF, HWP)", type=["docx", "pdf", "hwp"])

    if uploaded_file is not None:
        file_type = uploaded_file.name.split(".")[-1].lower()
        
        if file_type == "docx":
            doc = Document(uploaded_file)
            content = "\n".join([p.text for p in doc.paragraphs])
            st.text_area("DOCX 내용", content, height=600)
        
        elif file_type == "pdf":
            pdf = PdfReader(uploaded_file)
            content = ""
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    content += text + "\n"
            st.text_area("PDF 내용", content, height=600)
        
        else:
            st.warning("지원되지 않는 파일 형식입니다.")


    else:
        st.write('SOP 문서를 업로드 해주세요')
# 고칠부분 하이라이팅해주기(db 연결, 청크 id 연결해서 불러오기)
# 하이라이팅 부분 클릭하면 색깔바뀌면서 유사도 높은 gmp들 불러오기    
# 
#         
elif choice == "과거 기록":
    st.session_state.what = "과거 기록"  
    st.write('너의 과거입니다')


